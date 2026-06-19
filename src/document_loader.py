# src/document_loader.py

import os

from docx import Document

from langchain.schema import Document as LangchainDocument
from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter


def split_documents(documents):

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len,
        separators=["\n\n", "\n", " ", ""]
    )

    chunks = text_splitter.split_documents(documents)

    print(f"✅ Split into {len(chunks)} chunks")

    if chunks:

        print("\n📌 Sample chunk preview:")
        print(
            f"Content: "
            f"{chunks[0].page_content[:200]}..."
        )
        print(
            f"Metadata: "
            f"{chunks[0].metadata}"
        )

    return chunks


def load_and_split_pdf(file_path: str):

    if not os.path.exists(file_path):
        raise FileNotFoundError(
            f"File not found: {file_path}"
        )

    print(f"📄 Loading PDF: {file_path}")

    loader = PyPDFLoader(file_path)

    pages = loader.load()

    print(f"✅ Loaded {len(pages)} pages")

    return split_documents(pages)


def load_and_split_docx(file_path: str):

    if not os.path.exists(file_path):
        raise FileNotFoundError(
            f"File not found: {file_path}"
        )

    print(f"📄 Loading DOCX: {file_path}")

    doc = Document(file_path)

    text = "\n".join(
        para.text
        for para in doc.paragraphs
    )

    docs = [
        LangchainDocument(
            page_content=text,
            metadata={
                "source": file_path
            }
        )
    ]

    return split_documents(docs)


def load_and_split_text(file_path: str):

    if not os.path.exists(file_path):
        raise FileNotFoundError(
            f"File not found: {file_path}"
        )

    print(f"📄 Loading Text File: {file_path}")

    with open(
        file_path,
        "r",
        encoding="utf-8"
    ) as f:

        text = f.read()

    docs = [
        LangchainDocument(
            page_content=text,
            metadata={
                "source": file_path
            }
        )
    ]

    return split_documents(docs)


def load_document(file_path: str):

    extension = os.path.splitext(
        file_path
    )[1].lower()

    if extension == ".pdf":
        return load_and_split_pdf(
            file_path
        )

    elif extension == ".docx":
        return load_and_split_docx(
            file_path
        )

    elif extension in [
        ".txt",
        ".md"
    ]:
        return load_and_split_text(
            file_path
        )

    else:
        raise ValueError(
            f"Unsupported file type: {extension}"
        )


def get_document_info(chunks: list):

    total_chars = sum(
        len(chunk.page_content)
        for chunk in chunks
    )

    pages = max(
        (
            chunk.metadata.get("page", 0)
            for chunk in chunks
        ),
        default=0
    ) + 1

    return {
        "total_chunks": len(chunks),
        "total_characters": total_chars,
        "avg_chunk_size":
            total_chars // len(chunks)
            if chunks else 0,
        "pages": pages
    }